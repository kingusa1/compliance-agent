"""Extract Phase-2 compliance docs (.docx + .pdf) into plain markdown.

Outputs to .planning/phase2-docs/<slug>.md so the orchestrator and any
subagent can read the content without having to re-run docx/pdf parsers
each time. Idempotent — overwrites on re-run.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document  # python-docx
from PyPDF2 import PdfReader

ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "compliance-docs"
OUT = ROOT / ".planning" / "phase2-docs"
OUT.mkdir(parents=True, exist_ok=True)


def slug(name: str) -> str:
    return (
        name.lower()
        .replace(" ", "_")
        .replace(".docx", "")
        .replace(".pdf", "")
        .replace("(", "")
        .replace(")", "")
        .replace(",", "")
        .replace("'", "")
    )


def docx_to_md(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = [f"# {path.stem}\n"]
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        style = (p.style.name or "").lower()
        if style.startswith("heading 1"):
            parts.append(f"\n## {p.text}\n")
        elif style.startswith("heading 2"):
            parts.append(f"\n### {p.text}\n")
        elif style.startswith("heading"):
            parts.append(f"\n#### {p.text}\n")
        else:
            parts.append(p.text)
    # Tables
    for ti, table in enumerate(doc.tables):
        parts.append(f"\n\n### Table {ti + 1}\n")
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            parts.append(rows[0])
            parts.append(header_sep)
            parts.extend(rows[1:])
    return "\n".join(parts)


def pdf_to_md(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = [f"# {path.stem}\n"]
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        parts.append(f"\n## Page {i + 1}\n")
        parts.append(text)
    return "\n".join(parts)


def main() -> int:
    if not SRC.exists():
        print(f"missing source: {SRC}", file=sys.stderr)
        return 1
    written = 0
    for p in sorted(SRC.rglob("*")):
        if not p.is_file():
            continue
        ext = p.suffix.lower()
        try:
            if ext == ".docx":
                md = docx_to_md(p)
            elif ext == ".pdf":
                md = pdf_to_md(p)
            else:
                continue
        except Exception as e:  # noqa: BLE001
            print(f"FAIL  {p.name}: {type(e).__name__}: {e}", file=sys.stderr)
            continue
        rel = p.relative_to(SRC)
        out_name = slug(str(rel).replace("\\", "/").replace("/", "__")) + ".md"
        (OUT / out_name).write_text(md, encoding="utf-8")
        print(f"OK    {p.name} -> {out_name} ({len(md)} chars)")
        written += 1
    print(f"\nwrote {written} files to {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
